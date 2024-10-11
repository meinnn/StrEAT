from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import sys

# Word 문서 템플릿 로드
doc = Document('/truck.docx')

ownerName = sys.argv[1]
gender = sys.argv[2]
age = sys.argv[3]
birth = sys.argv[4]
address = sys.argv[5]
home_num = sys.argv[6]
phone_num = sys.argv[7]
email = sys.argv[8]
sns = sys.argv[9]
truckName = sys.argv[10]
businessNum = sys.argv[11]
eventName = sys.argv[12]

# 문서 내 모든 표를 순회하며 각 셀의 자리표시자를 대체
for table in doc.tables:
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if '대표자' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(ownerName)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '성 별' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(gender)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '연 령' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(age + '세')
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '주민등록번호' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(birth)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '주 소' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(address)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '연 락 처' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    if '일반전화' in row.cells[i+1].text:
                        row.cells[i + 1].text = ''
                        cell_to_update = row.cells[i + 1].paragraphs[0].add_run(
                            home_num)
                        cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if '휴대폰' in row.cells[i + 1].text:
                        row.cells[i + 1].text = ''
                        cell_to_update = row.cells[i + 1].paragraphs[0].add_run(
                            phone_num)
                        cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif 'ON-LINE' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    if 'E-Mail' in row.cells[i + 1].text:
                        row.cells[i + 1].text = ''
                        cell_to_update = row.cells[i + 1].paragraphs[0].add_run(
                            email)
                        cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    if 'SNS' in row.cells[i + 1].text:
                        row.cells[i + 1].text = ''
                        cell_to_update = row.cells[i + 1].paragraphs[0].add_run(
                            'Intagram@: ' + sns)
                        cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '푸드트럭명' in cell.text:
                if i + 1 < len(row.cells):  # 오른쪽 셀이 존재하는지 확인
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(truckName)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '사업자등록\n유무' in cell.text:
                if i + 1 < len(row.cells):
                    if '사업자등록' in row.cells[i + 1].text:
                        row.cells[i + 1].text = ''
                        cell_to_update = row.cells[i + 1].paragraphs[0].add_run('■ 사업자등록')
                        cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '사업자등록번호' in cell.text:
                if i + 1 < len(row.cells):
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(businessNum)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif '신청행사명' in cell.text:
                if i + 1 < len(row.cells):
                    row.cells[i + 1].text = ''
                    cell_to_update = row.cells[i + 1].paragraphs[0].add_run(eventName)
                    cell_to_update.font.highlight_color = WD_COLOR_INDEX.YELLOW

# 결과 문서 저장
doc.save('/입점신청서.docx')
